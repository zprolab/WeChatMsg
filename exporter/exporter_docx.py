import os
import time
import traceback
from re import findall

import docx
from docx import shared
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

from wxManager import Me, MessageType
from exporter.exporter import ExporterBase, get_new_filename
from wxManager.decrypt.decrypt_dat import decode_dat
from wxManager.log import logger
from wxManager.model import QuoteMessage, LinkMessage

# 要删除的编码字符
encoded_chars = b'\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f'

# 创建一个字典，将要删除的字符映射为 None
char_mapping = {char: None for char in encoded_chars}


def filter_control_characters(input_string):
    """
    过滤掉不可打印字符
    @param input_string:
    @return:
    """

    # 过滤掉非可打印字符
    filtered_string = input_string.translate(char_mapping)

    return filtered_string


class DocxExporter(ExporterBase):
    def __init__(
            self,
            database,
            contact,
            output_dir,
            type_,  # 导出文件类型
            message_types: set[MessageType] = None,  # 导出的消息类型
            time_range=None,  # 导出的日期范围
            group_members: set[str] = None,  # 群聊中只导出这些人的聊天记录
            progress_callback=None,  # 进度回调函数，func(progress:float)
            finish_callback=None,  # 导出完成回调函数
            msg_num_per_docx=500  # 每个docx文档的消息数量
    ):
        super().__init__(database, contact, output_dir, type_, message_types, time_range, group_members,
                         progress_callback, finish_callback)  # 调用父类的构造函数
        self.msg_num_per_docx = msg_num_per_docx

    def add_text_in(self, paragraph, content):
        try:
            paragraph.add_run(content)
        except ValueError:
            try:
                str_content = filter_control_characters(content)
                paragraph.add_run(str_content)
            except ValueError:
                logger.error(f'非法字符:{content}')
                paragraph.add_run('非法字符')

    def add_text_message(self, doc, avatar_path, is_send, content, display_name=''):
        content_cell = self.create_table(doc, is_send, avatar_path)
        if display_name:
            self.add_text_in(content_cell.paragraphs[0], display_name + '\n')
        self.add_text_in(content_cell.paragraphs[0], content)
        if is_send:
            p = content_cell.paragraphs[0]
        doc.add_paragraph()

    def text(self, doc, message):
        avatar = self.get_avatar_path(message, True)
        self.add_text_message(doc, avatar, message.is_sender, message.content,
                              message.display_name if self.contact.is_chatroom() else '')

    def image(self, doc, message):
        is_send = message.is_sender
        avatar = self.get_avatar_path(message, True)
        content = self.create_table(doc, is_send, avatar)
        if self.contact.is_chatroom():
            content.paragraphs[0].add_run(message.display_name + '\n')
        message.set_file_name()
        image_dir = os.path.join(self.origin_path, 'image')
        image_path = decode_dat(
            Me().xor_key,
            os.path.join(Me().wx_dir, message.path),
            os.path.join(image_dir, message.str_time[:7]),
            message.file_name
        )
        if image_path and os.path.exists(image_path):
            try:
                run = content.paragraphs[0].add_run()
                run.add_picture(image_path, height=shared.Inches(2))
                doc.add_paragraph()
            except Exception:
                try:
                    # 有些jpg图片的元数据丢失，垂直dpi丢失导致除0异常，需要手动指定图片长宽
                    run = content.paragraphs[0].add_run()
                    run.add_picture(image_path, height=shared.Inches(2), width=shared.Inches(2))
                    doc.add_paragraph()
                except:
                    print("Error!image")
                    logger.error(image_path)
                    logger.error(traceback.format_exc())
        else:
            content.paragraphs[0].add_run('【图片丢失】')

    def refermsg(self, doc, message: QuoteMessage):
        """
        处理回复消息
        @param doc:
        @param message:
        @return:
        """
        is_send = message.is_sender
        if message.quote_message.type == MessageType.Quote:
            refer_msg = f'{message.quote_message.display_name}:{message.quote_message.content}'
        else:
            refer_msg = f'{message.quote_message.display_name}:{message.quote_message.to_text()}'
        avatar = self.get_avatar_path(message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        self.add_text_in(content_cell.paragraphs[0], message.content)
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        reply_p = content_cell.add_paragraph()
        self.add_text_in(reply_p, refer_msg)
        run = reply_p.runs[0]
        '''设置被回复内容格式'''
        run.font.color.rgb = shared.RGBColor(121, 121, 121)
        run.font_size = shared.Inches(0.3)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

        if is_send:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            reply_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def set_table_different_width(self, table, widths):
        """表格分别设置列宽，单位为Cm"""
        for x, width in enumerate(widths):
            for cell in table.columns[x].cells:
                cell.width = Inches(width)

    def link(self, doc, message: LinkMessage):
        """
        处理回复消息
        @param doc:
        @param message:
        @return:
        """
        is_send = message.is_sender

        avatar = self.get_avatar_path(message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        # 创建一个包含两行两列的表格
        # 第一行用于放置标题、内容以及缩略图
        # 第二行合并两列，用于显示应用名
        table = content_cell.add_table(rows=3, cols=2)
        self.set_table_different_width(table, [10086, 0.5])
        # 第一行左侧单元格：添加标题和内容
        cell_header = table.cell(0, 0)
        cell_header.merge(table.cell(0, 1))
        # 添加标题
        p_title = cell_header.paragraphs[0]
        self.add_text_in(p_title, message.title)
        run = p_title.runs[0]
        run.font.size = Pt(12)  # 设置字体大小
        if message.href:
            r_id = p_title.part.relate_to(message.href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)  # 关联超链接

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            run.font.underline = True
            hyperlink.append(run._r)

            p_title._element.append(hyperlink)

        # 添加内容
        cell_content = table.cell(1, 0)
        self.add_text_in(cell_content.paragraphs[0], message.description)

        # 第一行右侧单元格：添加缩略图
        cell_right = table.cell(1, 1)
        cell_right.width = shared.Inches(0.5)
        p_img = cell_right.paragraphs[0]
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中显示
        # 插入图片，调整图片宽度（需确保thumbnail.jpg图片存在）
        # p_img.add_run().add_picture(
        #     r"E:\Project\Python\MemoTrace\dist\MemoTrace-2.1.0\data\聊天记录\涵涵(wxid_9snu8vlykdpt22)\涵涵_14\word\media\image1.png",
        #     width=Inches(0.5))

        # 第二行：合并两个单元格显示应用名
        cell_app = table.cell(2, 0)
        cell_app.merge(table.cell(2, 1))
        self.add_text_in(cell_app.paragraphs[0], message.app_name)
        cell_app.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 应用名居中
        self.delete_paragraph(content_cell.paragraphs[0])
        doc.add_paragraph()

    def add_system_text(self, doc, text):
        if not text:
            return
        p = doc.add_paragraph(text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        run.font.color.rgb = shared.RGBColor(121, 121, 121)
        run.font.size = Pt(9)
        # run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    def system_msg(self, doc, message):
        str_content = message.content
        str_content = str_content.replace('<![CDATA[', "").replace(
            ' <a href="weixin://revoke_edit_click">重新编辑</a>]]>', "")
        res = findall('(</{0,1}(img|revo|_wc_cus|a).*?>)', str_content)
        for xmlstr, b in res:
            str_content = str_content.replace(xmlstr, "")
        self.add_system_text(doc, str_content)

    def delete_paragraph(self, paragraph):
        """删除某一段落"""
        p = paragraph._element
        tc = p.getparent()
        tc.remove(p)
        p._p = None
        p._element = None

    def create_table(self, doc, is_send, avatar_path):
        """
        #! 创建一个1*2表格
        #! isSend = 1 (0,0)存聊天内容，(0,1)存头像
        #! isSend = 0 (0,0)存头像，(0,1)存聊天内容
        #! 返回聊天内容的坐标
        """
        table = doc.add_table(rows=1, cols=2, style='Normal Table')
        table.cell(0, 1).height = shared.Inches(0.5)
        table.cell(0, 0).height = shared.Inches(0.5)
        if is_send:
            '''表格右对齐'''
            table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            avatar = table.cell(0, 1).paragraphs[0].add_run()
            '''插入头像，设置头像宽度'''
            try:
                avatar.add_picture(avatar_path, width=shared.Inches(0.5))
            except:
                logger.error('头像插入失败')
            '''设置单元格宽度跟头像一致'''
            table.cell(0, 1).width = shared.Inches(0.5)
            content_cell = table.cell(0, 0)
            '''聊天内容右对齐'''
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            avatar = table.cell(0, 0).paragraphs[0].add_run()
            try:
                avatar.add_picture(avatar_path, width=shared.Inches(0.5))
            except:
                logger.error('头像插入失败')
            '''设置单元格宽度'''
            table.cell(0, 0).width = shared.Inches(0.5)
            content_cell = table.cell(0, 1)
        '''聊天内容垂直居中对齐'''
        content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return content_cell

    def export(self):
        print(f"【开始导出 DOCX {self.contact.remark}】")
        origin_path = self.origin_path
        messages = self.database.get_messages(self.contact.wxid, time_range=self.time_range)
        total_steps = len(messages)
        self.save_avatars()

        def newdoc():
            nonlocal docx_num, doc
            doc = docx.Document()
            doc.styles["Normal"].font.name = "Cambria"
            doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            core_properties = doc.core_properties
            core_properties.author = 'MemoTrace'  # 作者
            core_properties.comments = 'generated by MemoTrace'  # 注释
            docx_num += 1

        doc = None
        docx_num = 0
        newdoc()
        selected_msg_cnt = 0
        for index, message in enumerate(messages):
            if index and index % 1000 == 0:
                self.update_progress_callback(index / total_steps)
            if not self.is_selected(message):
                continue

            if not self._is_running:
                break
            selected_msg_cnt += 1
            type_ = message.type
            timestamp = message.timestamp
            if self.is_5_min(timestamp):
                str_time = message.str_time
                self.add_system_text(doc, str_time)
            if type_ == MessageType.System:
                self.system_msg(doc, message)
            elif type_ == MessageType.Quote:
                self.refermsg(doc, message)
            elif type_ == MessageType.Image:
                self.image(doc, message)
            elif type_ in {MessageType.LinkMessage, MessageType.Applet, MessageType.Music}:
                self.link(doc, message)
            else:
                try:
                    avatar = self.get_avatar_path(message, True)
                    self.add_text_message(doc, avatar, message.is_sender, message.to_text(),
                                          message.display_name if self.contact.is_chatroom() else '')
                except:
                    pass
            if selected_msg_cnt % self.msg_num_per_docx == 0 or index == total_steps - 1:
                filename = os.path.join(origin_path, f"{self.contact.remark}_{docx_num}.docx")
                filename = get_new_filename(filename)
                try:
                    doc.save(filename)
                except PermissionError:
                    filename = os.path.join(origin_path, f"{self.contact.remark}_{docx_num}_{str(time.time())}.docx")
                    doc.save(filename)
                except:
                    pass
                newdoc()
        self.update_progress_callback(1)
        print(f"【完成导出 DOCX {self.contact.remark}】")
        self.finish_callback(self.exporter_id)
